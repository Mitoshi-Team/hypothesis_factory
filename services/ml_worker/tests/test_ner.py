from __future__ import annotations

import sqlite3
import sys
import tempfile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from models import EntityLabel, SourceType, UnifiedDocument, UnifiedElement
from ner.db_handler import DBHandler
from ner.extractors import (
    DBExtractor,
    DocExtractor,
    ExcelExtractor,
    MarkdownExtractor,
    PDFExtractor,
    TextExtractor,
    get_extractor,
)
from ner.ner_extractor import NER_LABEL_MAP, NERExtractor
from ner.router import EXTENSION_MAP, route_file

# ─── Fixtures ────────────────────────────────────────────────────────────────


@pytest.fixture
def tmp_dir():
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


@pytest.fixture
def txt_file(tmp_dir: Path) -> Path:
    path = tmp_dir / "test.txt"
    path.write_text(
        "Ниобий повышает жаропрочность стали.\n"
        "The addition of nickel increases corrosion resistance.\n"
        "Содержание хрома в нержавеющей стали составляет 18%.\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def md_file(tmp_dir: Path) -> Path:
    path = tmp_dir / "test.md"
    path.write_text(
        "# Introduction\n"
        "This document describes material properties.\n"
        "\n"
        "- Ниобий\n"
        "- Хром\n"
        "- Никель\n"
        "\n"
        "## Properties\n"
        "Жаропрочность and corrosion resistance are key.\n",
        encoding="utf-8",
    )
    return path


@pytest.fixture
def pdf_file(tmp_dir: Path) -> Path:
    path = tmp_dir / "test.pdf"
    _create_minimal_pdf(path, "Test PDF content with nickel alloy")
    return path


@pytest.fixture
def xlsx_file(tmp_dir: Path) -> Path:
    import pandas as pd

    path = tmp_dir / "test.xlsx"
    df = pd.DataFrame(
        {
            "Material": ["Ниобий", "Хром", "Никель"],
            "Property": ["Жаропрочность", "Коррозия", "Прочность"],
            "Value": [1200, 800, 950],
        }
    )
    df.to_excel(str(path), sheet_name="Materials", index=False)
    return path


@pytest.fixture
def docx_file(tmp_dir: Path) -> Path:
    from docx import Document

    path = tmp_dir / "test.docx"
    doc = Document()
    doc.add_heading("Materials Report", level=1)
    doc.add_paragraph("This report describes nickel alloy properties.")
    doc.add_paragraph(
        "Ниобий повышает жаропрочность стали.", style="List Bullet"
    )
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Material"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Ниобий"
    table.cell(1, 1).text = "1200"
    table.cell(2, 0).text = "Хром"
    table.cell(2, 1).text = "800"
    doc.save(str(path))
    return path


@pytest.fixture
def db_file(tmp_dir: Path) -> Path:
    path = tmp_dir / "test.db"
    conn = sqlite3.connect(str(path))
    conn.execute(
        "CREATE TABLE materials (id INTEGER, name TEXT, property TEXT, value REAL)"
    )
    conn.execute(
        "INSERT INTO materials VALUES (1, 'Ниобий', 'Жаропрочность', 1200)"
    )
    conn.execute("INSERT INTO materials VALUES (2, 'Хром', 'Коррозия', 800)")
    conn.commit()
    conn.close()
    return path


def _create_minimal_pdf(path: Path, text: str):
    text_bytes = text.encode("latin-1")
    stream_data = b"BT /F1 12 Tf 100 700 Td(" + text_bytes + b")Tj ET"

    objs = [
        b"%PDF-1.4",
        b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj",
        b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj",
        b"3 0 obj<</Type/Page/Parent 2 0 R/MediaBox[0 0 612 792]/Contents 4 0 R"
        b"/Resources<</Font<</F1 5 0 R>>>>>>endobj",
        b"4 0 obj<</Length "
        + str(len(stream_data)).encode()
        + b">>\nstream\n"
        + stream_data
        + b"\nendstream\nendobj",
        b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj",
    ]

    body = b"\n".join(objs)
    offsets = [0] * 6
    offsets[1] = 0
    current = len(objs[1]) + 1
    for i in range(2, 6):
        offsets[i] = current
        current += len(objs[i]) + 1

    xref_offset = current
    xref_lines = [b"xref", b"0 6", b"0000000000 65535 f "]
    for i in range(1, 6):
        xref_lines.append(f"{offsets[i]:010d} 00000 n ".encode())
    xref = b"\n".join(xref_lines)
    trailer = (
        b"\ntrailer\n<</Size 6/Root 1 0 R>>\nstartxref\n"
        + str(xref_offset).encode()
        + b"\n%%EOF"
    )
    path.write_bytes(body + b"\n" + xref + trailer)


def _make_mock_gliner(entities: list[dict]):
    class MockGLiNER:
        def __init__(self, ents):
            self._ents = ents

        def predict_entities(
            self, text: str, labels: list[str], threshold: float = 0.5
        ):
            return self._ents

    return MockGLiNER(entities)


# ─── Router Tests ────────────────────────────────────────────────────────────


class TestRouter:
    def test_route_txt(self):
        assert route_file("file.txt") == SourceType.TEXT

    def test_route_markdown(self):
        assert route_file("file.md") == SourceType.MARKDOWN
        assert route_file("file.markdown") == SourceType.MARKDOWN
        assert route_file("file.mdown") == SourceType.MARKDOWN

    def test_route_pdf(self):
        assert route_file("file.pdf") == SourceType.PDF

    def test_route_excel(self):
        assert route_file("file.xlsx") == SourceType.EXCEL
        assert route_file("file.xls") == SourceType.EXCEL

    def test_route_word(self):
        assert route_file("file.docx") == SourceType.WORD
        assert route_file("file.doc") == SourceType.WORD

    def test_route_database(self):
        assert route_file("file.db") == SourceType.DATABASE
        assert route_file("file.sqlite") == SourceType.DATABASE
        assert route_file("file.sql") == SourceType.DATABASE
        assert route_file("file.csv") == SourceType.DATABASE

    def test_route_case_insensitive(self):
        assert route_file("FILE.TXT") == SourceType.TEXT
        assert route_file("File.PDF") == SourceType.PDF

    def test_route_unknown_raises(self):
        with pytest.raises(ValueError, match="Unsupported file extension"):
            route_file("file.xyz")

    def test_route_no_extension_raises(self):
        with pytest.raises(ValueError):
            route_file("file")

    def test_extension_map_coverage(self):
        all_types = set(SourceType)
        mapped_types = set(EXTENSION_MAP.values())
        assert all_types == mapped_types, (
            f"Missing: {all_types - mapped_types}"
        )


# ─── Extractor Tests ─────────────────────────────────────────────────────────


class TestGetExtractor:
    def test_get_extractor_returns_correct_type(self):
        assert isinstance(get_extractor(SourceType.TEXT), TextExtractor)
        assert isinstance(
            get_extractor(SourceType.MARKDOWN), MarkdownExtractor
        )
        assert isinstance(get_extractor(SourceType.PDF), PDFExtractor)
        assert isinstance(get_extractor(SourceType.EXCEL), ExcelExtractor)
        assert isinstance(get_extractor(SourceType.DATABASE), DBExtractor)
        assert isinstance(get_extractor(SourceType.WORD), DocExtractor)

    def test_get_extractor_unknown_raises(self):
        with pytest.raises(ValueError):
            get_extractor("unknown")


class TestTextExtractor:
    def test_extract_returns_unified_document(self, txt_file: Path):
        extractor = TextExtractor()
        doc = extractor.extract(str(txt_file))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.TEXT
        assert doc.source_uri == str(txt_file.absolute())

    def test_extract_contains_elements(self, txt_file: Path):
        extractor = TextExtractor()
        doc = extractor.extract(str(txt_file))
        assert len(doc.elements) > 0
        assert any("ниобий" in el.text.lower() for el in doc.elements)

    def test_extract_has_checksum(self, txt_file: Path):
        extractor = TextExtractor()
        doc = extractor.extract(str(txt_file))
        assert doc.checksum is not None
        assert len(doc.checksum) == 64

    def test_extract_empty_file(self, tmp_dir: Path):
        path = tmp_dir / "empty.txt"
        path.write_text("")
        extractor = TextExtractor()
        doc = extractor.extract(str(path))
        assert len(doc.elements) == 0


class TestMarkdownExtractor:
    def test_extract_parses_headings(self, md_file: Path):
        extractor = MarkdownExtractor()
        doc = extractor.extract(str(md_file))
        assert len(doc.elements) > 0

    def test_extract_returns_unified_document(self, md_file: Path):
        extractor = MarkdownExtractor()
        doc = extractor.extract(str(md_file))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.MARKDOWN


class TestPDFExtractor:
    def test_extract_fallback_returns_document(self, pdf_file: Path):
        extractor = PDFExtractor()
        doc = extractor.extract(str(pdf_file))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.PDF

    def test_extract_fallback_contains_text(self, pdf_file: Path):
        extractor = PDFExtractor()
        doc = extractor.extract(str(pdf_file))
        texts = [el.text for el in doc.elements]
        combined = " ".join(texts)
        assert "nickel" in combined.lower() or "alloy" in combined.lower()

    def test_extract_nonexistent_file_raises(self):
        extractor = PDFExtractor()
        with pytest.raises(FileNotFoundError):
            extractor.extract("/nonexistent/file.pdf")


class TestExcelExtractor:
    def test_extract_returns_document(self, xlsx_file: Path):
        extractor = ExcelExtractor()
        doc = extractor.extract(str(xlsx_file))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.EXCEL

    def test_extract_contains_table_data(self, xlsx_file: Path):
        extractor = ExcelExtractor()
        doc = extractor.extract(str(xlsx_file))
        assert len(doc.elements) > 0
        table_el = doc.elements[0]
        assert table_el.table_data is not None
        assert "Ниобий" in str(table_el.table_data.rows)

    def test_extract_sheet_name_in_table_data(self, xlsx_file: Path):
        extractor = ExcelExtractor()
        doc = extractor.extract(str(xlsx_file))
        td = doc.elements[0].table_data
        assert td.name == "Materials"


class TestDocExtractor:
    def test_extract_returns_document(self, docx_file: Path):
        extractor = DocExtractor()
        doc = extractor.extract(str(docx_file))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.WORD

    def test_extract_contains_elements(self, docx_file: Path):
        extractor = DocExtractor()
        doc = extractor.extract(str(docx_file))
        assert len(doc.elements) > 0

    def test_extract_parses_heading(self, docx_file: Path):
        extractor = DocExtractor()
        doc = extractor.extract(str(docx_file))
        titles = [el for el in doc.elements if el.type.value == "title"]
        assert any("materials" in el.text.lower() for el in titles)

    def test_extract_contains_table(self, docx_file: Path):
        extractor = DocExtractor()
        doc = extractor.extract(str(docx_file))
        tables = [el for el in doc.elements if el.table_data is not None]
        assert len(tables) > 0
        td = tables[0].table_data
        assert td.columns == ["Material", "Value"]

    def test_extract_contains_text(self, docx_file: Path):
        extractor = DocExtractor()
        doc = extractor.extract(str(docx_file))
        texts = [el.text for el in doc.elements]
        combined = " ".join(texts)
        assert "nickel" in combined.lower()

    def test_extract_docx_with_cyrillic(self, docx_file: Path):
        extractor = DocExtractor()
        doc = extractor.extract(str(docx_file))
        texts = " ".join(el.text for el in doc.elements)
        assert "ниобий" in texts.lower()

    def test_extract_doc_unsupported_fallback(self, tmp_dir: Path):
        path = tmp_dir / "old.doc"
        path.write_text("dummy", encoding="utf-8")
        extractor = DocExtractor()
        doc = extractor.extract(str(path))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.WORD


class TestDBExtractor:
    def test_extract_returns_document(self, db_file: Path):
        extractor = DBExtractor()
        doc = extractor.extract(str(db_file))
        assert isinstance(doc, UnifiedDocument)
        assert doc.source_type == SourceType.DATABASE

    def test_extract_reads_tables(self, db_file: Path):
        extractor = DBExtractor()
        doc = extractor.extract(str(db_file))
        assert len(doc.elements) > 0
        for el in doc.elements:
            assert el.table_data is not None

    def test_extract_table_has_rows(self, db_file: Path):
        extractor = DBExtractor()
        doc = extractor.extract(str(db_file))
        td = doc.elements[0].table_data
        assert len(td.rows) >= 2
        assert td.name == "materials"

    def test_resolve_db_url_sqlite(self):
        url = DBExtractor._resolve_db_url("/tmp/test.db")  # noqa: S108
        assert url.startswith("sqlite:///")

    def test_resolve_db_url_csv_raises(self):
        with pytest.raises(ValueError, match="CSV"):
            DBExtractor._resolve_db_url("/tmp/test.csv")  # noqa: S108

    def test_extract_nonexistent_db_raises(self):
        extractor = DBExtractor()
        with pytest.raises(Exception):
            extractor.extract("/nonexistent/test.db")


# ─── NERExtractor Tests ──────────────────────────────────────────────────────


class TestNERExtractor:
    def test_extract_entities_with_mocked_gliner(self):
        extractor = NERExtractor(model_name="mock-model")
        extractor._model = _make_mock_gliner(
            [
                {"text": "Ниобий", "label": "MATERIAL", "score": 0.95},
                {"text": "Chromium", "label": "MATERIAL", "score": 0.87},
            ]
        )

        elements = [
            UnifiedElement(
                type="text",
                text="Ниобий and Chromium are materials",
                source_type=SourceType.TEXT,
            )
        ]
        entities = extractor.extract_entities(elements)
        assert len(entities) == 2

        names = [e.name for e in entities]
        assert "ниобий" in names
        assert "chromium" in names

    def test_extract_entities_label_is_material(self):
        extractor = NERExtractor(model_name="mock-model")
        extractor._model = _make_mock_gliner(
            [
                {"text": "Niobium", "label": "MATERIAL", "score": 0.9},
            ]
        )

        elements = [
            UnifiedElement(
                type="text",
                text="Niobium is a metal",
                source_type=SourceType.TEXT,
            )
        ]
        entities = extractor.extract_entities(elements)
        assert len(entities) == 1
        assert entities[0].label == EntityLabel.MATERIAL

    def test_extract_entities_respects_label(self):
        extractor = NERExtractor(model_name="mock-model")
        extractor._model = _make_mock_gliner(
            [
                {"text": "отжиг", "label": "PROCESS", "score": 0.9},
                {"text": "1100 °C", "label": "PARAMETER", "score": 0.85},
            ]
        )

        elements = [
            UnifiedElement(
                type="text",
                text="отжиг при 1100 °C",
                source_type=SourceType.TEXT,
            )
        ]
        entities = extractor.extract_entities(elements)
        assert len(entities) == 2
        labels = {e.name: e.label for e in entities}
        assert labels["отжиг"] == EntityLabel.PROCESS
        assert labels["1100 °c"] == EntityLabel.PARAMETER

    def test_extract_entities_deduplicates(self):
        extractor = NERExtractor(model_name="mock-model")
        extractor._model = _make_mock_gliner(
            [
                {"text": "Ниобий", "label": "MATERIAL", "score": 0.95},
                {"text": "Ниобий", "label": "MATERIAL", "score": 0.90},
            ]
        )

        elements = [
            UnifiedElement(
                type="text", text="Ниобий Ниобий", source_type=SourceType.TEXT
            )
        ]
        entities = extractor.extract_entities(elements)
        assert len(entities) == 1

    def test_extract_entities_empty_elements(self):
        extractor = NERExtractor(model_name="mock-model")
        entities = extractor.extract_entities([])
        assert entities == []

    def test_extract_entities_empty_text(self):
        extractor = NERExtractor(model_name="mock-model")
        elements = [
            UnifiedElement(type="text", text="", source_type=SourceType.TEXT)
        ]
        entities = extractor.extract_entities(elements)
        assert entities == []

    def test_normalize_name(self):
        extractor = NERExtractor(model_name="mock")
        assert extractor._normalize_name(" Niobium ") == "niobium"
        assert extractor._normalize_name(" 1100 °C ") == "1100 °c"
        assert extractor._normalize_name("Сталь 20Х23Н18") == "сталь 20х23н18"


class TestNERLabelMap:
    def test_gliner_label_map(self):
        expected = {
            "MATERIAL": EntityLabel.MATERIAL,
            "PROCESS": EntityLabel.PROCESS,
            "PROPERTY": EntityLabel.PROPERTY,
            "PARAMETER": EntityLabel.PARAMETER,
        }
        assert NER_LABEL_MAP == expected


# ─── DBHandler Tests ─────────────────────────────────────────────────────────


class TestDBHandler:
    def test_copy_tables_logs(self):
        handler = DBHandler()
        doc = UnifiedDocument(
            source_type=SourceType.TEXT, source_uri="/dev/null"
        )
        handler.copy_tables(doc)

    def test_save_entities_logs(self):
        handler = DBHandler()
        handler.save_entities([])

    def test_save_relations_logs(self):
        handler = DBHandler()
        handler.save_relations([])
