from __future__ import annotations

import hashlib
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Optional

from models import (
    ElementType,
    SourceType,
    TableCell,
    TableData,
    UnifiedDocument,
    UnifiedElement,
)


class BaseExtractor(ABC):
    source_type: SourceType

    @abstractmethod
    def extract(self, file_path: str) -> UnifiedDocument:
        ...

    def _make_document(
        self,
        file_path: str,
        elements: list[UnifiedElement],
        title: Optional[str] = None,
        language: str = "en",
        extractor: str = "unknown",
    ) -> UnifiedDocument:
        path = Path(file_path)
        mime_map: dict[str, str] = {
            ".txt": "text/plain",
            ".md": "text/markdown",
            ".markdown": "text/markdown",
            ".pdf": "application/pdf",
            ".xls": "application/vnd.ms-excel",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".csv": "text/csv",
            ".sql": "application/sql",
            ".db": "application/octet-stream",
            ".sqlite": "application/octet-stream",
            ".doc": "application/msword",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
        source_bytes = path.read_bytes()

        doc = UnifiedDocument(
            source_type=self.source_type,
            source_uri=str(path.absolute()),
            mime_type=mime_map.get(path.suffix.lower()),
            title=title or path.stem,
            elements=elements,
            language=language,
            extractor=extractor,
        )
        doc.compute_checksum(source_bytes)
        return doc


class PDFExtractor(BaseExtractor):
    source_type = SourceType.PDF

    def extract(self, file_path: str) -> UnifiedDocument:
        try:
            return self._extract_with_docling(file_path)
        except ImportError:
            return self._fallback_extract(file_path)

    def _extract_with_docling(self, file_path: str) -> UnifiedDocument:
        from docling.document_converter import DocumentConverter

        converter = DocumentConverter()
        docling_doc = converter.convert(file_path).document

        elements: list[UnifiedElement] = []
        for item, level in docling_doc.iterate_items():
            el_type = self._map_docling_type(item)
            text = item.text if hasattr(item, "text") else ""
            element = UnifiedElement(
                type=el_type,
                text=text or "",
                level=level,
                page=getattr(item, "page", None),
                extractor="docling",
                source_type=SourceType.PDF,
            )

            if el_type == ElementType.TABLE:
                table_data = self._convert_docling_table(item)
                if table_data:
                    element.table_data = table_data

            elements.append(element)

        title = Path(file_path).stem
        return self._make_document(
            file_path=file_path,
            elements=elements,
            title=title,
            extractor="docling",
        )

    def _map_docling_type(self, item: Any) -> ElementType:
        type_str = type(item).__name__.lower() if hasattr(item, "__class__") else ""
        mapping = {
            "title": ElementType.TITLE,
            "heading": ElementType.TITLE,
            "paragraph": ElementType.TEXT,
            "text": ElementType.TEXT,
            "list": ElementType.LIST_ITEM,
            "table": ElementType.TABLE,
            "figure": ElementType.IMAGE,
            "picture": ElementType.IMAGE,
            "formula": ElementType.FORMULA,
            "code": ElementType.CODE,
            "caption": ElementType.CAPTION,
            "header": ElementType.HEADER,
            "footer": ElementType.FOOTER,
            "reference": ElementType.REFERENCE,
        }
        for key, val in mapping.items():
            if key in type_str:
                return val
        return ElementType.TEXT

    def _convert_docling_table(self, item: Any) -> Optional[TableData]:
        try:
            import pandas as pd

            df: pd.DataFrame = item.export_to_dataframe()
            cells = []
            for row_idx, (_, row) in enumerate(df.iterrows()):
                for col_idx, value in enumerate(row):
                    cells.append(
                        TableCell(
                            row=row_idx,
                            col=col_idx,
                            text=str(value) if value is not None else "",
                            is_header=row_idx == 0,
                        )
                    )
            return TableData(
                columns=list(df.columns),
                rows=[list(row) for _, row in df.iterrows()],
                cells=cells,
                markdown=item.export_to_markdown() if hasattr(item, "export_to_markdown") else None,
            )
        except Exception:
            return None

    def _fallback_extract(self, file_path: str) -> UnifiedDocument:
        from PyPDF2 import PdfReader

        reader = PdfReader(file_path)
        elements: list[UnifiedElement] = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                elements.append(
                    UnifiedElement(
                        type=ElementType.TEXT,
                        text=text,
                        page=page_num + 1,
                        extractor="pypdf2",
                        source_type=SourceType.PDF,
                    )
                )
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="pypdf2",
        )


class TextExtractor(BaseExtractor):
    source_type = SourceType.TEXT

    def extract(self, file_path: str) -> UnifiedDocument:
        try:
            return self._extract_with_unstructured(file_path)
        except ImportError:
            return self._fallback_extract(file_path)

    def _extract_with_unstructured(self, file_path: str) -> UnifiedDocument:
        from unstructured.partition.text import partition_text

        raw_elements = partition_text(filename=file_path)
        elements = [self._convert_unstructured_element(el) for el in raw_elements]
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="unstructured",
        )

    def _fallback_extract(self, file_path: str) -> UnifiedDocument:
        raw = Path(file_path).read_bytes()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("cp1251", errors="replace")
        elements = []
        if text.strip():
            elements.append(
                UnifiedElement(
                    type=ElementType.TEXT,
                    text=text,
                    extractor="builtin",
                    source_type=SourceType.TEXT,
                )
            )
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="builtin",
        )

    def _convert_unstructured_element(self, el: Any) -> UnifiedElement:
        category_map = {
            "Title": ElementType.TITLE,
            "Header": ElementType.HEADER,
            "Footer": ElementType.FOOTER,
            "ListItem": ElementType.LIST_ITEM,
            "Table": ElementType.TABLE,
            "FigureCaption": ElementType.CAPTION,
            "Caption": ElementType.CAPTION,
            "Formula": ElementType.FORMULA,
            "CodeSnippet": ElementType.CODE,
            "PageBreak": ElementType.TEXT,
            "UncategorizedText": ElementType.TEXT,
            "NarrativeText": ElementType.TEXT,
        }
        category = el.category if hasattr(el, "category") else "NarrativeText"
        el_type = category_map.get(category, ElementType.TEXT)
        text = el.text if hasattr(el, "text") else str(el)
        return UnifiedElement(
            type=el_type,
            text=text or "",
            extractor="unstructured",
            source_type=self.source_type,
        )


class MarkdownExtractor(BaseExtractor):
    source_type = SourceType.MARKDOWN

    def extract(self, file_path: str) -> UnifiedDocument:
        try:
            return self._extract_with_unstructured(file_path)
        except ImportError:
            return self._fallback_extract(file_path)

    def _extract_with_unstructured(self, file_path: str) -> UnifiedDocument:
        from unstructured.partition.md import partition_md

        raw_elements = partition_md(filename=file_path)
        elements = [self._convert_unstructured_element(el) for el in raw_elements]
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="unstructured",
        )

    def _fallback_extract(self, file_path: str) -> UnifiedDocument:
        raw = Path(file_path).read_bytes()
        try:
            text = raw.decode("utf-8")
        except UnicodeDecodeError:
            text = raw.decode("cp1251", errors="replace")
        elements: list[UnifiedElement] = []
        for line in text.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                elements.append(
                    UnifiedElement(
                        type=ElementType.TITLE,
                        text=stripped.lstrip("#").strip(),
                        level=level,
                        extractor="builtin",
                        source_type=SourceType.MARKDOWN,
                    )
                )
            elif stripped.startswith("-") or stripped.startswith("*"):
                elements.append(
                    UnifiedElement(
                        type=ElementType.LIST_ITEM,
                        text=stripped.lstrip("-*").strip(),
                        extractor="builtin",
                        source_type=SourceType.MARKDOWN,
                    )
                )
            else:
                elements.append(
                    UnifiedElement(
                        type=ElementType.TEXT,
                        text=stripped,
                        extractor="builtin",
                        source_type=SourceType.MARKDOWN,
                    )
                )
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="builtin",
        )

    def _convert_unstructured_element(self, el: Any) -> UnifiedElement:
        category_map = {
            "Title": ElementType.TITLE,
            "Header": ElementType.HEADER,
            "Footer": ElementType.FOOTER,
            "ListItem": ElementType.LIST_ITEM,
            "Table": ElementType.TABLE,
            "FigureCaption": ElementType.CAPTION,
            "Caption": ElementType.CAPTION,
            "Formula": ElementType.FORMULA,
            "CodeSnippet": ElementType.CODE,
            "PageBreak": ElementType.TEXT,
            "UncategorizedText": ElementType.TEXT,
            "NarrativeText": ElementType.TEXT,
        }
        category = el.category if hasattr(el, "category") else "NarrativeText"
        el_type = category_map.get(category, ElementType.TEXT)
        text = el.text if hasattr(el, "text") else str(el)
        return UnifiedElement(
            type=el_type,
            text=text or "",
            extractor="unstructured",
            source_type=self.source_type,
        )


class ExcelExtractor(BaseExtractor):
    source_type = SourceType.EXCEL

    def extract(self, file_path: str) -> UnifiedDocument:
        import pandas as pd

        elements: list[UnifiedElement] = []
        xls = pd.ExcelFile(file_path)
        try:
            sheet_names = xls.sheet_names
        finally:
            xls.close()

        for sheet_name in sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            cells = [
                TableCell(
                    row=row_idx,
                    col=col_idx,
                    text=str(value) if pd.notna(value) else "",
                    is_header=row_idx == 0,
                )
                for row_idx, (_, row) in enumerate(df.iterrows())
                for col_idx, value in enumerate(row)
            ]
            cols = [str(c) for c in df.columns]
            rows = [
                [str(v) if pd.notna(v) else None for v in row]
                for _, row in df.iterrows()
            ]
            try:
                markdown = df.to_markdown(index=False)
            except ImportError:
                markdown = None
            table_data = TableData(
                name=sheet_name,
                columns=cols,
                rows=rows,
                cells=cells,
                markdown=markdown,
            )
            elements.append(
                UnifiedElement(
                    type=ElementType.TABLE,
                    text=f"Sheet: {sheet_name}\n{df.to_string(index=False)}",
                    table_data=table_data,
                    extractor="pandas",
                    source_type=SourceType.EXCEL,
                )
            )

        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="pandas",
        )


class DBExtractor(BaseExtractor):
    source_type = SourceType.DATABASE

    def extract(self, file_path: str) -> UnifiedDocument:
        from sqlalchemy import create_engine, inspect, text

        db_url = self._resolve_db_url(file_path)
        engine = create_engine(db_url)
        inspector = inspect(engine)
        elements: list[UnifiedElement] = []

        for table_name in inspector.get_table_names():
            columns = [col["name"] for col in inspector.get_columns(table_name)]
            with engine.connect() as conn:
                result = conn.execute(text(f"SELECT * FROM {table_name}"))
                rows = [list(row) for row in result.fetchall()]

            rows_serializable = [
                [str(v) if v is not None else None for v in row] for row in rows
            ]
            cells = [
                TableCell(
                    row=row_idx,
                    col=col_idx,
                    text=str(value) if value is not None else "",
                    is_header=row_idx == 0,
                )
                for row_idx, row in enumerate(rows)
                for col_idx, value in enumerate(row)
            ]

            table_data = TableData(
                name=table_name,
                columns=columns,
                rows=rows_serializable,
                cells=cells,
            )
            elements.append(
                UnifiedElement(
                    type=ElementType.TABLE,
                    text=f"Table: {table_name}\nColumns: {', '.join(columns)}\nRows: {len(rows)}",
                    table_data=table_data,
                    extractor="sqlalchemy",
                    source_type=SourceType.DATABASE,
                )
            )

        engine.dispose()
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="sqlalchemy",
        )

    @staticmethod
    def _resolve_db_url(file_path: str) -> str:
        path = Path(file_path)
        if path.suffix.lower() in (".db", ".sqlite", ".sqlite3"):
            return f"sqlite:///{path.absolute()}"
        if path.suffix.lower() == ".csv":
            raise ValueError(
                "CSV files are not directly supported as databases. "
                "Use ExcelExtractor or import CSV into a database first."
            )
        return file_path


class DocExtractor(BaseExtractor):
    source_type = SourceType.WORD

    def extract(self, file_path: str) -> UnifiedDocument:
        ext = Path(file_path).suffix.lower()
        if ext == ".doc":
            return self._extract_doc(file_path)
        return self._extract_docx(file_path)

    def _extract_docx(self, file_path: str) -> UnifiedDocument:
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        doc = DocxDocument(file_path)
        elements: list[UnifiedElement] = []
        table_iter = iter(doc.tables)

        for child in doc.element.body:
            tag = child.tag
            if tag.endswith("tbl"):
                table = next(table_iter, None)
                if table is not None:
                    td = self._convert_docx_table(table)
                    elements.append(
                        UnifiedElement(
                            type=ElementType.TABLE,
                            text=td.markdown or "",
                            table_data=td,
                            extractor="python-docx",
                            source_type=SourceType.WORD,
                        )
                    )
            elif tag.endswith("p"):
                from docx.text.paragraph import Paragraph

                para = Paragraph(child, doc)
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name.lower() if para.style else ""
                if "heading" in style_name or "title" in style_name:
                    level = None
                    for part in style_name.replace("heading", "").strip().split():
                        if part.isdigit():
                            level = int(part)
                            break
                    elements.append(
                        UnifiedElement(
                            type=ElementType.TITLE,
                            text=text,
                            level=level or 1,
                            extractor="python-docx",
                            source_type=SourceType.WORD,
                        )
                    )
                elif style_name.startswith("list"):
                    elements.append(
                        UnifiedElement(
                            type=ElementType.LIST_ITEM,
                            text=text,
                            extractor="python-docx",
                            source_type=SourceType.WORD,
                        )
                    )
                else:
                    elements.append(
                        UnifiedElement(
                            type=ElementType.TEXT,
                            text=text,
                            extractor="python-docx",
                            source_type=SourceType.WORD,
                        )
                    )

        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="python-docx",
        )

    def _extract_doc(self, file_path: str) -> UnifiedDocument:
        try:
            return self._extract_with_unstructured(file_path)
        except ImportError:
            return self._make_document(
                file_path=file_path,
                elements=[],
                extractor="builtin",
            )

    def _extract_with_unstructured(self, file_path: str) -> UnifiedDocument:
        from unstructured.partition.doc import partition_doc

        raw_elements = partition_doc(filename=file_path)
        elements = [self._convert_unstructured_element(el) for el in raw_elements]
        return self._make_document(
            file_path=file_path,
            elements=elements,
            extractor="unstructured",
        )

    @staticmethod
    def _convert_unstructured_element(el: Any) -> UnifiedElement:
        category_map = {
            "Title": ElementType.TITLE,
            "Header": ElementType.HEADER,
            "Footer": ElementType.FOOTER,
            "ListItem": ElementType.LIST_ITEM,
            "Table": ElementType.TABLE,
            "FigureCaption": ElementType.CAPTION,
            "Caption": ElementType.CAPTION,
            "Formula": ElementType.FORMULA,
            "CodeSnippet": ElementType.CODE,
            "PageBreak": ElementType.TEXT,
            "UncategorizedText": ElementType.TEXT,
            "NarrativeText": ElementType.TEXT,
        }
        category = el.category if hasattr(el, "category") else "NarrativeText"
        el_type = category_map.get(category, ElementType.TEXT)
        text = el.text if hasattr(el, "text") else str(el)
        return UnifiedElement(
            type=el_type,
            text=text or "",
            extractor="unstructured",
            source_type=SourceType.WORD,
        )

    @staticmethod
    def _convert_docx_table(table: Any) -> TableData:
        cells: list[TableCell] = []
        rows_data: list[list[str | None]] = []
        columns: list[str] = []

        for row_idx, row in enumerate(table.rows):
            row_values: list[str | None] = []
            for col_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                row_values.append(text or None)
                cells.append(
                    TableCell(
                        row=row_idx,
                        col=col_idx,
                        text=text,
                        is_header=row_idx == 0,
                    )
                )
            rows_data.append(row_values)
            if row_idx == 0:
                columns = [str(c) for c in row_values]

        return TableData(
            columns=columns,
            rows=rows_data,
            cells=cells,
        )


EXTRACTOR_MAP: dict[SourceType, type[BaseExtractor]] = {
    SourceType.PDF: PDFExtractor,
    SourceType.TEXT: TextExtractor,
    SourceType.MARKDOWN: MarkdownExtractor,
    SourceType.EXCEL: ExcelExtractor,
    SourceType.DATABASE: DBExtractor,
    SourceType.WORD: DocExtractor,
}


def get_extractor(source_type: SourceType) -> BaseExtractor:
    extractor_cls = EXTRACTOR_MAP.get(source_type)
    if extractor_cls is None:
        raise ValueError(f"No extractor for source type: {source_type}")
    return extractor_cls()
